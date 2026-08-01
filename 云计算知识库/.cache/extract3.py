# -*- coding: utf-8 -*-
import os, re, json, sys

CACHE = r"C:\Users\liang\ObsidianVault\云计算知识库\.cache"
BASE = r"E:\云原生大礼包（资料）\11.PDF书籍&面试题\电子书"

# key -> (相对 BASE 的文件名 或 绝对路径, type)
SRC = {
    # ---- 顶层 PDF ----
    "docker-book":        ("Docker书.pdf", "pdf"),
    "docker-shizhan":     ("Docker实战.pdf", "pdf"),
    "docker-tujie":       ("Docker实战(图解）.pdf", "pdf"),
    "docker-jingdian":    ("Docker经典实例.pdf", "pdf"),
    "git-team":           ("Git团队协作.pdf", "pdf"),
    "go-note":            ("Go 学习笔记 第四版.pdf", "pdf"),
    "go-action":          ("Go语言实战 .pdf", "pdf"),
    "k8s-zhinan":         ("Kubernetes指南.pdf", "pdf"),
    "linux-ops":          ("Linux系统架构与运维实战.pdf", "pdf"),
    "bigdata-alibaba":    ("《大数据之路：阿里巴巴大数据实践》.pdf", "pdf"),
    "docker-k8s-sanjia":  ("从Docker→kubernetes(三架马车）.pdf", "pdf"),
    "monitoring-docx":    ("常用的运维监控系统.docx", "docx"),
    # ---- 云原生迷你电子书合集 (去重 k8s-zhinan) ----
    "container-cases":    (r"云原生迷你电子书合集\10个精选的容器应用案例.pdf", "pdf"),
    "devops-cloud":       (r"云原生迷你电子书合集\DevOps在智能云时代的开发与交付.pdf", "pdf"),
    "docker-rumen":       (r"云原生迷你电子书合集\Docker从入门到实践.pdf", "pdf"),
    "envoy-doc":          (r"云原生迷你电子书合集\ENVOY 官方文档.pdf", "pdf"),
    "istio-doc":          (r"云原生迷你电子书合集\Istio1.6官方文档中文版.pdf", "pdf"),
    "k8s-paas":           (r"云原生迷你电子书合集\Kubernetes PaaS冲击波.pdf", "pdf"),
    "k8s-practice":       (r"云原生迷你电子书合集\Kubernetes 实践指南（Kubernetes Practice Guide）.pdf", "pdf"),
    "spring-ioc":         (r"云原生迷你电子书合集\Spring的IoC容器.pdf", "pdf"),
    "ceph-doc-zh":        (r"云原生迷你电子书合集\ceph详细中文文档.pdf", "pdf"),
    "vmware-vsphere":     (r"云原生迷你电子书合集\《 VMware vSphere with Kubernetes 基础知识》白皮书：云原生领域的一大进步.pdf", "pdf"),
    "cloudnative-trend":  (r"云原生迷你电子书合集\云原生时代容器云的技术发展趋势.pdf", "pdf"),
    "cloudnative-practice":(r"云原生迷你电子书合集\云原生的技术探索与落地实践.pdf", "pdf"),
    "microservice-devops":(r"云原生迷你电子书合集\微服务与DevOps技术内参.pdf", "pdf"),
    "k8s-deploy-10":      (r"云原生迷你电子书合集\有关 Kubernetes 部署的10 个注意事项.pdf", "pdf"),
    # ---- 249个Shell脚本 ----
    "shell-249":          (r"249个Shell脚本\【强推】249个拿来即用shell脚本.pdf", "pdf"),
    "shell-rumen-note":   (r"249个Shell脚本\Shell脚本编程入门笔记.pdf", "pdf"),
    "shell-base":         (r"249个Shell脚本\Shell编程基础.pdf", "pdf"),
    "shell-spec":         (r"249个Shell脚本\shell编写规范.pdf", "pdf"),
    "shell-master":       (r"249个Shell脚本\Shell从入门到精通.pdf", "pdf"),
    "shell-base2":        (r"249个Shell脚本\Shell脚本基础.pdf", "pdf"),
    "shell-100cases":     (r"249个Shell脚本\100个Linux+Shell脚本经典案例.pdf", "pdf"),
    "shell-strategy-cn":  (r"249个Shell脚本\LINUX SHELL脚本攻略(中文版带书签).pdf", "pdf"),
    "shell-strategy":     (r"249个Shell脚本\Linux Shell脚本攻略.pdf", "pdf"),
    "shell-strategy-2nd": (r"249个Shell脚本\Linux Shell脚本攻略（第2版）.pdf", "pdf"),
    "shell-linuxshell":   (r"249个Shell脚本\LinuxShell.pdf", "pdf"),
    "shell-study":        (r"249个Shell脚本\Shell脚本学习指南.pdf", "pdf"),
    "shell-pro-master":   (r"249个Shell脚本\shell编程从入门到精通.pdf", "pdf"),
    "sed-principle":      (r"249个Shell脚本\玩透sed：探究sed原理.pdf", "pdf"),
    "regex-master":       (r"249个Shell脚本\精通正则表达式.pdf", "pdf"),
}

def sanitize(s):
    s = re.sub(r'[\\/:*?"<>|]', '_', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s[:60]

def extract_pdf(key, path):
    import fitz
    d = os.path.join(CACHE, key)
    os.makedirs(d, exist_ok=True)
    doc = fitz.open(path)
    full = []
    pages = doc.page_count
    for pno in range(pages):
        t = doc.load_page(pno).get_text("text")
        full.append(t)
    full_text = "\n".join(full)
    with open(os.path.join(d, "full.txt"), "w", encoding="utf-8") as f:
        f.write(full_text)
    # TOC -> chapters
    toc = doc.get_toc()
    chapters = []
    if toc:
        # top-level entries (level 1) define chapter boundaries
        tops = [(lvl, title, page) for (lvl, title, page) in toc if lvl == 1]
        if len(tops) >= 3:
            for i, (lvl, title, page0) in enumerate(tops):
                page1 = tops[i+1][2] if i+1 < len(tops) else pages + 1
                p0 = max(1, page0); p1 = max(p0, page1)
                txt = "\n".join(full[p0-1:p1-1])
                if not txt.strip():
                    continue
                fname = f"ch{str(i+1).zfill(2)}_{sanitize(title)}.txt"
                with open(os.path.join(d, fname), "w", encoding="utf-8") as f:
                    f.write(txt)
                chapters.append({"file": fname, "title": title, "pages": p1-p0, "chars": len(txt)})
    doc.close()
    avg = (len(full_text) / pages) if pages else 0
    manifest = {"key": key, "pages": pages, "chapters": chapters,
                "avg_bytes_per_page": round(avg, 1),
                "type": "pdf",
                "quality": "TEXT" if avg > 700 else ("MIXED" if avg > 150 else "SCANNED")}
    json.dump(manifest, open(os.path.join(d, "manifest.json"), "w", encoding="utf-8"),
              ensure_ascii=False, indent=2)
    print(f"[OK] {key:20s} pages={pages:5d} ch={len(chapters):3d} avgB/pg={avg:7.1f} {manifest['quality']}")

def extract_docx(key, path):
    import docx
    d = os.path.join(CACHE, key)
    os.makedirs(d, exist_ok=True)
    doc = docx.Document(path)
    paras = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paras)
    with open(os.path.join(d, "full.txt"), "w", encoding="utf-8") as f:
        f.write(full_text)
    # split by headings (level 1/2)
    chapters = []
    buf = []
    cur_title = "正文"
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip():
            if buf:
                txt = "\n".join(buf)
                if txt.strip():
                    chapters.append({"file": f"ch{str(len(chapters)+1).zfill(2)}.txt",
                                     "title": cur_title, "pages": 0, "chars": len(txt)})
            cur_title = p.text.strip()
            buf = [p.text]
        else:
            buf.append(p.text)
    if buf:
        txt = "\n".join(buf)
        if txt.strip():
            chapters.append({"file": f"ch{str(len(chapters)+1).zfill(2)}.txt",
                             "title": cur_title, "pages": 0, "chars": len(txt)})
    # write chapter files
    ci = 0
    cur_title = "正文"; buf = []
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip():
            if buf:
                txt = "\n".join(buf)
                if txt.strip():
                    ci += 1
                    with open(os.path.join(d, f"ch{str(ci).zfill(2)}.txt"), "w", encoding="utf-8") as f:
                        f.write(txt)
            cur_title = p.text.strip(); buf = [p.text]
        else:
            buf.append(p.text)
    if buf:
        txt = "\n".join(buf)
        if txt.strip():
            ci += 1
            with open(os.path.join(d, f"ch{str(ci).zfill(2)}.txt"), "w", encoding="utf-8") as f:
                f.write(txt)
    manifest = {"key": key, "pages": len(paras), "chapters": chapters,
                "avg_bytes_per_page": round(len(full_text)/max(1,len(paras)),1),
                "type": "docx", "quality": "TEXT"}
    json.dump(manifest, open(os.path.join(d, "manifest.json"), "w", encoding="utf-8"),
              ensure_ascii=False, indent=2)
    print(f"[OK] {key:20s} paras={len(paras):5d} ch={len(chapters):3d} docx")

def main():
    force = "--force" in sys.argv
    done = 0; skip = 0
    for key, (rel, typ) in SRC.items():
        path = rel if os.path.isabs(rel) else os.path.join(BASE, rel)
        if not os.path.exists(path):
            print(f"[MISS] {key}: {path}")
            continue
        out_full = os.path.join(CACHE, key, "full.txt")
        if os.path.exists(out_full) and not force:
            skip += 1
            continue
        try:
            if typ == "pdf":
                extract_pdf(key, path)
            else:
                extract_docx(key, path)
            done += 1
        except Exception as e:
            print(f"[ERR] {key}: {e}")
    print(f"\n=== 完成：本次提取 {done} 个，跳过(已存在) {skip} 个，共 {len(SRC)} 个 ===")

if __name__ == "__main__":
    main()
